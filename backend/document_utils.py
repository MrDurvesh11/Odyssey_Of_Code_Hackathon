import os
from typing import Dict, List, Any, Optional, Tuple

class DocumentExtractor:
    """Utility class for extracting text from various document formats"""
    
    @staticmethod
    def extract_text(file_path: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.txt':
            return DocumentExtractor._extract_from_txt(file_path)
        elif file_extension == '.pdf':
            return DocumentExtractor._extract_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return DocumentExtractor._extract_from_word(file_path)
        else:
            return "", f"Unsupported file format: {file_extension}"
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from a text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read(), None
        except Exception as e:
            return "", f"Error extracting text from {file_path}: {str(e)}"
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from a PDF file"""
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text, None
        except ImportError:
            return "", "PyPDF2 is not installed. Install it using: pip install PyPDF2"
        except Exception as e:
            return "", f"Error extracting text from PDF {file_path}: {str(e)}"
    
    @staticmethod
    def _extract_from_word(file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from a Word document"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text, None
        except ImportError:
            return "", "python-docx is not installed. Install it using: pip install python-docx"
        except Exception as e:
            return "", f"Error extracting text from Word document {file_path}: {str(e)}"

class DocumentStructureAnalyzer:
    """Analyze document structure to improve chunking quality"""
    
    @staticmethod
    def identify_sections(text: str) -> Dict[str, str]:
        """
        Identify logical sections in a document
        
        Args:
            text: Document text
            
        Returns:
            Dictionary of section name to section content
        """
        # Simple section detection based on common patterns
        import re
        
        # Look for common section headers
        section_patterns = [
            r'^#+\s+(.+?)$',  # Markdown headers
            r'^(\d+(\.\d+)*)\s+(.+?)$',  # Numbered sections
            r'^([A-Z][A-Za-z\s]+:)$',  # Title followed by colon
            r'^([A-Z][A-Za-z\s]+)$'  # All caps or title case lines
        ]
        
        current_section = "Introduction"
        sections = {current_section: ""}
        
        for line in text.split('\n'):
            # Check if this line is a section header
            is_header = False
            for pattern in section_patterns:
                if re.match(pattern, line, re.MULTILINE):
                    current_section = line.strip()
                    if current_section not in sections:
                        sections[current_section] = ""
                    is_header = True
                    break
            
            if not is_header:
                sections[current_section] += line + "\n"
        
        return sections

def get_document_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from a document
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Dictionary of metadata
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path)
    
    metadata = {
        "file_name": file_name,
        "file_extension": file_extension,
        "file_size": file_size,
        "file_path": file_path
    }
    
    # Extract additional metadata based on file type
    if file_extension == '.pdf':
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)
                if pdf_reader.metadata:
                    metadata["author"] = pdf_reader.metadata.get('/Author')
                    metadata["creation_date"] = pdf_reader.metadata.get('/CreationDate')
                    metadata["producer"] = pdf_reader.metadata.get('/Producer')
        except:
            pass
    
    elif file_extension in ['.docx', '.doc']:
        try:
            import docx
            doc = docx.Document(file_path)
            metadata["paragraph_count"] = len(doc.paragraphs)
            core_properties = doc.core_properties
            metadata["author"] = core_properties.author
            metadata["created"] = core_properties.created
            metadata["modified"] = core_properties.modified
        except:
            pass
    
    return metadata
